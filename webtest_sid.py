from bs4 import BeautifulSoup
from bs4.element import Comment
from urllib.request import Request, urlopen
from docx import Document


url = 'https://www.creepypasta.com/if-youre-armed-and-at-the-glenmont-metro-please-shoot-me/'
text_file = 'sitedata.txt'
out_file = 'tester.rtf'

def tag_visible(element):
    #if element is in any of the following headers it returns false
    if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']: 
        return False
    #if element is a comment, ignore it as well (false)
    if isinstance(element, Comment):
        return False
    #Element cannot be any of those headers or a comment
    return True 


def text_from_html(body):
    soup = BeautifulSoup(body, 'html.parser')
    texts = soup.findAll(text=True) #This gets inside HTML tags to get the text
    visible_texts = filter(tag_visible, texts)  #Filters the text
    #good = (visible_texts)
    
    #return " ".join(t.strip() for t in visible_texts)
    #print(u" ".join(g.strip() for g in visible_texts))
    text_to_doc = []
    for g in visible_texts: #goes through all the fucking texts found
        if not g.isspace(): #filters out all the whitespace
            #print("/////////////////////")
            #print((g.strip()))
            
            with open (text_file, 'a') as goods:   #Appends to the file rather than overwriting
                goods.write("\n\n") #goods is just the file bitch, write two spaces to file
                try:
                    goods.write(str(g.strip()))
                    break_up(str(g.strip()))    #sends to break up string into list to bold it
                except UnicodeEncodeError:
                    #print("lol you failed:", str(g.strip()))
                    pass
                    #goods.write(str((g.strip()).encode('utf-8')))
    
def break_up(stuff):
    string = stuff
    x = string.split(" ")
    bold(x) #makes stuff bold
    #print("This is x: ",x)

def bold(stuff):
    temp = []
    out = open(out_file,'a')
    for words in stuff:
        first = '\033[1m' + words[:len(words)//2] + '\033[0m' #this bolds the first half of the word
        out.write("""\\b """ + words[:len(words)//2] + """\\b0""")
        last = words[len(words)//2:]    #last half of word
        out.write(words[len(words)//2:] + """ """)
        good_word = first + last        #Combines the bolded first half and the last half
        temp.append(good_word)          #puts the modified word into a list
    sentence = " ".join(temp)           #joins all the stuff in the list into a sentence
    out.write("""\line\line""")
    out.close()
    #print(sentence)                     #Outputs the sentence
    # send to word doc or sth
    word_doc(sentence)


def word_doc(bold):
    pass
    #doc = Document()
    #doc.add_paragraph(bold.encode('utf-8'))
    #doc.save("out.docx")

def init(): #supposed to spawn file, and make it blank for new data
    
    with open(text_file, 'w') as f:
        f.truncate(0)

    out = open(out_file,'w')
    out.write("""{\\rtf1""")
    out.close()
    #Gets contents of webpage and can bypass common bot detectors lol
    page = Request(url, headers={'User-Agent': 'Mozilla/5.0'})  
    #Reads html stuff from the contents
    html = urlopen(page).read()

    doc = Document()
    doc._body.clear_content()
    doc.save("out.docx")

    return html

def main():
    webpage_stuff = init() #Gets contents of webpage
    #html = urllib.request.urlopen('https://www.yahoo.com/').read()

    text_from_html(webpage_stuff) #Extracts text from the webpage contents hopefully
    out = open(out_file,'a')
    out.write("""}""")
    out.close()
    #word_doc()

main()
'''
with open('sitedata.txt', 'w') as f:
    f.write(str(stuff.encode("utf-8")))
'''